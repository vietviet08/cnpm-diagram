from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading('Đặc Tả Use Case', level=1)

use_cases = [
    {
        "Tiêu đề": "Đăng Ký Người Dùng",
        "Mô tả": "Use case này cho phép người dùng mới đăng ký tài khoản trong hệ thống.",
        "Tác nhân": "Người dùng",
        "Điều kiện kích hoạt": "Người dùng muốn tạo một tài khoản mới.",
        "Tiền điều kiện": "- Người dùng chưa đăng ký trong hệ thống.",
        "Hậu điều kiện": "- Một tài khoản người dùng mới được tạo và lưu trữ trong cơ sở dữ liệu.",
        "Luồng sự kiện chính": [
            "1. Người dùng truy cập vào trang đăng ký.",
            "2. Người dùng nhập các thông tin bao gồm tên đăng nhập, mật khẩu, và email.",
            "3. Người dùng gửi biểu mẫu đăng ký.",
            "4. Hệ thống xác thực các thông tin đã cung cấp.",
            "5. Nếu xác thực thành công, hệ thống tạo một tài khoản người dùng mới.",
            "6. Hệ thống hiển thị thông báo thành công cho người dùng."
        ],
        "Luồng sự kiện phụ": [
            "- Nếu xác thực thất bại, hệ thống hiển thị thông báo lỗi và yêu cầu người dùng sửa lại thông tin."
        ]
    },
    {
        "Tiêu đề": "Đăng Nhập Người Dùng",
        "Mô tả": "Use case này cho phép người dùng đã đăng ký đăng nhập vào tài khoản của họ.",
        "Tác nhân": "Người dùng",
        "Điều kiện kích hoạt": "Người dùng muốn đăng nhập vào tài khoản của họ.",
        "Tiền điều kiện": "- Người dùng đã đăng ký trong hệ thống.",
        "Hậu điều kiện": "- Người dùng đăng nhập vào tài khoản của họ.",
        "Luồng sự kiện chính": [
            "1. Người dùng truy cập vào trang đăng nhập.",
            "2. Người dùng nhập tên đăng nhập và mật khẩu.",
            "3. Người dùng gửi biểu mẫu đăng nhập.",
            "4. Hệ thống xác thực thông tin đăng nhập.",
            "5. Nếu xác thực thành công, hệ thống đăng nhập người dùng.",
            "6. Hệ thống hiển thị bảng điều khiển của người dùng."
        ],
        "Luồng sự kiện phụ": [
            "- Nếu xác thực thất bại, hệ thống hiển thị thông báo lỗi và yêu cầu người dùng sửa lại thông tin đăng nhập."
        ]
    },
    {
        "Tiêu đề": "Tạo Tin Tuyển Dụng",
        "Mô tả": "Use case này cho phép một công ty tạo một tin tuyển dụng mới.",
        "Tác nhân": "Công ty",
        "Điều kiện kích hoạt": "Công ty muốn tạo một tin tuyển dụng mới.",
        "Tiền điều kiện": "- Công ty đã đăng ký trong hệ thống và đã đăng nhập.",
        "Hậu điều kiện": "- Một tin tuyển dụng mới được tạo và lưu trữ trong cơ sở dữ liệu.",
        "Luồng sự kiện chính": [
            "1. Công ty truy cập vào trang tạo tin tuyển dụng.",
            "2. Công ty nhập các thông tin về công việc bao gồm tiêu đề, mô tả, địa điểm, và mức lương.",
            "3. Công ty gửi biểu mẫu tin tuyển dụng.",
            "4. Hệ thống xác thực các thông tin đã cung cấp.",
            "5. Nếu xác thực thành công, hệ thống tạo một tin tuyển dụng mới.",
            "6. Hệ thống hiển thị thông báo thành công cho công ty."
        ],
        "Luồng sự kiện phụ": [
            "- Nếu xác thực thất bại, hệ thống hiển thị thông báo lỗi và yêu cầu công ty sửa lại thông tin."
        ]
    },
    {
        "Tiêu đề": "Nộp Đơn Xin Việc",
        "Mô tả": "Use case này cho phép người dùng nộp đơn xin việc cho một tin tuyển dụng.",
        "Tác nhân": "Người dùng",
        "Điều kiện kích hoạt": "Người dùng muốn nộp đơn xin việc cho một tin tuyển dụng.",
        "Tiền điều kiện": "- Người dùng đã đăng ký trong hệ thống và đã đăng nhập.",
        "Hậu điều kiện": "- Đơn xin việc của người dùng được gửi và lưu trữ trong cơ sở dữ liệu.",
        "Luồng sự kiện chính": [
            "1. Người dùng truy cập vào trang tin tuyển dụng.",
            "2. Người dùng nhấn nút 'Nộp đơn'.",
            "3. Người dùng chọn CV của mình và gửi biểu mẫu đơn xin việc.",
            "4. Hệ thống xác thực các thông tin đã cung cấp.",
            "5. Nếu xác thực thành công, hệ thống tạo một đơn xin việc mới.",
            "6. Hệ thống hiển thị thông báo thành công cho người dùng."
        ],
        "Luồng sự kiện phụ": [
            "- Nếu xác thực thất bại, hệ thống hiển thị thông báo lỗi và yêu cầu người dùng sửa lại thông tin."
        ]
    },
    {
        "Tiêu đề": "Gửi CV",
        "Mô tả": "Use case này cho phép người dùng gửi CV của họ.",
        "Tác nhân": "Người dùng",
        "Điều kiện kích hoạt": "Người dùng muốn gửi CV của họ.",
        "Tiền điều kiện": "- Người dùng đã đăng ký trong hệ thống và đã đăng nhập.",
        "Hậu điều kiện": "- CV của người dùng được gửi và lưu trữ trong cơ sở dữ liệu.",
        "Luồng sự kiện chính": [
            "1. Người dùng truy cập vào trang gửi CV.",
            "2. Người dùng tải lên tệp CV của mình.",
            "3. Người dùng gửi biểu mẫu CV.",
            "4. Hệ thống xác thực các thông tin đã cung cấp.",
            "5. Nếu xác thực thành công, hệ thống tạo một mục CV mới.",
            "6. Hệ thống hiển thị thông báo thành công cho người dùng."
        ],
        "Luồng sự kiện phụ": [
            "- Nếu xác thực thất bại, hệ thống hiển thị thông báo lỗi và yêu cầu người dùng sửa lại thông tin."
        ]
    },
    {
        "Tiêu đề": "Xem Tin Tuyển Dụng",
        "Mô tả": "Use case này cho phép người dùng xem các tin tuyển dụng.",
        "Tác nhân": "Người dùng",
        "Điều kiện kích hoạt": "Người dùng muốn xem các tin tuyển dụng.",
        "Tiền điều kiện": "- Người dùng đã đăng ký trong hệ thống và đã đăng nhập.",
        "Hậu điều kiện": "- Người dùng xem các tin tuyển dụng có sẵn.",
        "Luồng sự kiện chính": [
            "1. Người dùng truy cập vào trang tin tuyển dụng.",
            "2. Hệ thống hiển thị các tin tuyển dụng có sẵn.",
            "3. Người dùng chọn một tin tuyển dụng để xem chi tiết."
        ],
        "Luồng sự kiện phụ": [
            "- Không có"
        ]
    },
    {
        "Tiêu đề": "Đánh Giá Công Ty",
        "Mô tả": "Use case này cho phép người dùng gửi đánh giá cho một công ty.",
        "Tác nhân": "Người dùng",
        "Điều kiện kích hoạt": "Người dùng muốn gửi đánh giá cho một công ty.",
        "Tiền điều kiện": "- Người dùng đã đăng ký trong hệ thống và đã đăng nhập.",
        "Hậu điều kiện": "- Đánh giá của người dùng được gửi và lưu trữ trong cơ sở dữ liệu.",
        "Luồng sự kiện chính": [
            "1. Người dùng truy cập vào trang đánh giá công ty.",
            "2. Người dùng nhập các chi tiết đánh giá bao gồm đánh giá và nhận xét.",
            "3. Người dùng gửi biểu mẫu đánh giá.",
            "4. Hệ thống xác thực các thông tin đã cung cấp.",
            "5. Nếu xác thực thành công, hệ thống tạo một mục đánh giá mới.",
            "6. Hệ thống hiển thị thông báo thành công cho người dùng."
        ],
        "Luồng sự kiện phụ": [
            "- Nếu xác thực thất bại, hệ thống hiển thị thông báo lỗi và yêu cầu người dùng sửa lại thông tin."
        ]
    },
    {
        "Tiêu đề": "Tạo Nhóm Chat",
        "Mô tả": "Use case này cho phép người dùng tạo một nhóm chat mới.",
        "Tác nhân": "Người dùng",
        "Điều kiện kích hoạt": "Người dùng muốn tạo một nhóm chat mới.",
        "Tiền điều kiện": "- Người dùng đã đăng ký trong hệ thống và đã đăng nhập.",
        "Hậu điều kiện": "- Một nhóm chat mới được tạo và lưu trữ trong cơ sở dữ liệu.",
        "Luồng sự kiện chính": [
            "1. Người dùng truy cập vào trang tạo nhóm chat.",
            "2. Người dùng nhập các thông tin về nhóm chat bao gồm tên và mô tả.",
            "3. Người dùng gửi biểu mẫu nhóm chat.",
            "4. Hệ thống xác thực các thông tin đã cung cấp.",
            "5. Nếu xác thực thành công, hệ thống tạo một mục nhóm chat mới.",
            "6. Hệ thống hiển thị thông báo thành công cho người dùng."
        ],
        "Luồng sự kiện phụ": [
            "- Nếu xác thực thất bại, hệ thống hiển thị thông báo lỗi và yêu cầu người dùng sửa lại thông tin."
        ]
    },
    {
        "Tiêu đề": "Gửi Tin Nhắn",
        "Mô tả": "Use case này cho phép người dùng gửi tin nhắn trong một nhóm chat.",
        "Tác nhân": "Người dùng",
        "Điều kiện kích hoạt": "Người dùng muốn gửi tin nhắn trong một nhóm chat.",
        "Tiền điều kiện": "- Người dùng đã đăng ký trong hệ thống và đã đăng nhập.",
        "Hậu điều kiện": "- Tin nhắn của người dùng được gửi và lưu trữ trong cơ sở dữ liệu.",
        "Luồng sự kiện chính": [
            "1. Người dùng truy cập vào trang nhóm chat.",
            "2. Người dùng nhập tin nhắn của mình.",
            "3. Người dùng gửi biểu mẫu tin nhắn.",
            "4. Hệ thống xác thực các thông tin đã cung cấp.",
            "5. Nếu xác thực thành công, hệ thống tạo một mục tin nhắn mới.",
            "6. Hệ thống hiển thị tin nhắn trong nhóm chat."
        ],
        "Luồng sự kiện phụ": [
            "- Nếu xác thực thất bại, hệ thống hiển thị thông báo lỗi và yêu cầu người dùng sửa lại thông tin."
        ]
    },
    {
        "Tiêu đề": "Xem Tin Nhắn",
        "Mô tả": "Use case này cho phép người dùng xem các tin nhắn trong một nhóm chat.",
        "Tác nhân": "Người dùng",
        "Điều kiện kích hoạt": "Người dùng muốn xem các tin nhắn trong một nhóm chat.",
        "Tiền điều kiện": "- Người dùng đã đăng ký trong hệ thống và đã đăng nhập.",
        "Hậu điều kiện": "- Người dùng xem các tin nhắn trong nhóm chat.",
        "Luồng sự kiện chính": [
            "1. Người dùng truy cập vào trang nhóm chat.",
            "2. Hệ thống hiển thị các tin nhắn trong nhóm chat.",
            "3. Người dùng cuộn qua các tin nhắn để xem chúng."
        ],
        "Luồng sự kiện phụ": [
            "- Không có"
        ]
    }
]

for use_case in use_cases:
    doc.add_heading(use_case["Tiêu đề"], level=2)
    table = doc.add_table(rows=0, cols=2)
    
    for key, value in use_case.items():
        if key.startswith("Luồng sự kiện chính") or key.startswith("Luồng sự kiện phụ"):
            row = table.add_row().cells
            row[0].text = key
            row[1].text = '\n'.join(value)
        else:
            row = table.add_row().cells
            row[0].text = key
            row[1].text = value
    
    doc.add_paragraph()

doc.save("usecase_specifications_vietnamese.docx")
